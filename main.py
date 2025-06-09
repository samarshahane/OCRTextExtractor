from fastapi import FastAPI, File, UploadFile, Request
from fastapi.responses import HTMLResponse, FileResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from pathlib import Path
import shutil
import pdfplumber
import easyocr
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import cv2
import numpy as np

app = FastAPI()
templates = Jinja2Templates(directory="templates")
app.mount("/static", StaticFiles(directory="static"), name="static")

upload_folder = "uploads"
output_file = "output/formatted.docx"
Path(upload_folder).mkdir(exist_ok=True)
Path("output").mkdir(exist_ok=True)

reader = easyocr.Reader(['en'], gpu=False)
extracted_text = ""

@app.get("/", response_class=HTMLResponse)
async def index(request: Request):
    return templates.TemplateResponse("index.html", {"request": request, "preview_text": None, "show_download": False})

def save_file(file: UploadFile):
    file_path = Path(upload_folder) / file.filename
    with open(file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
    return str(file_path)

def align_paragraph(para, x0, page_width):
    margin = 50  # tighter margin
    center_range = (page_width * 0.4, page_width * 0.6)
    if x0 < center_range[0]:
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif x0 > center_range[1]:
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

def process_pdf(file_path):
    global extracted_text
    doc = Document()
    extracted_text = ""

    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_width = page.width
            # Handle tables first
            tables = page.extract_tables()
            if tables:
                for table in tables:
                    t = doc.add_table(rows=0, cols=len(table[0]))
                    for row in table:
                        row_cells = t.add_row().cells
                        for i, cell in enumerate(row):
                            row_cells[i].text = (cell or "").strip()
                            row_cells[i].paragraphs[0].runs[0].font.size = Pt(11)
                            row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                            extracted_text += (cell or "") + " | "
                        extracted_text += "\n"
                doc.add_paragraph()  # spacing between content

            # Then handle individual characters
            lines = page.extract_words(x_tolerance=1, y_tolerance=3, keep_blank_chars=True, use_text_flow=True)
            for word in lines:
                text = word['text']
                x0 = word['x0']
                size = round(float(word.get('size', 11)))
                para = doc.add_paragraph()
                run = para.add_run(text)
                run.font.size = Pt(size)
                align_paragraph(para, x0, page_width)
                extracted_text += text + " "

    doc.save(output_file)

def process_image(file_path):
    global extracted_text
    doc = Document()
    results = reader.readtext(file_path, detail=1)
    img = cv2.imread(file_path)
    width = img.shape[1]
    extracted_text = ""

    for (bbox, text, conf) in results:
        top_left = bbox[0]
        bottom_right = bbox[2]
        height = bottom_right[1] - top_left[1]
        x_pos = (top_left[0] + bottom_right[0]) / 2

        para = doc.add_paragraph()
        run = para.add_run(text)
        run.font.size = Pt(max(10, int(height / 2)))

        if x_pos < width * 0.4:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif x_pos > width * 0.6:
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        extracted_text += text + "\n"

    doc.save(output_file)

@app.post("/upload")
async def upload_file(request: Request, file: UploadFile = File(...)):
    file_path = save_file(file)
    ext = Path(file.filename).suffix.lower()

    if ext == ".pdf":
        process_pdf(file_path)
    elif ext in [".jpg", ".jpeg", ".png"]:
        process_image(file_path)
    else:
        return {"error": "Unsupported file type"}

    return templates.TemplateResponse("index.html", {
        "request": request,
        "preview_text": extracted_text,
        "show_download": True
    })

@app.get("/download")
async def download_docx():
    return FileResponse(output_file, media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document', filename="formatted.docx")
